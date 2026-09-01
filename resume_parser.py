"""
resume_parser.py
Extracts raw text from uploaded PDF or DOCX resumes.
"""

import io

import docx  # python-docx
import fitz  # PyMuPDF


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using PyMuPDF."""
    chunks = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            chunks.append(page.get_text())
    return "\n".join(chunks).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX (paragraphs + tables) using python-docx."""
    document = docx.Document(io.BytesIO(file_bytes))

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts).strip()


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Dispatch to the right extractor based on file extension."""
    ext = filename.lower().rsplit(".", 1)[-1]

    if ext == "pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext == "docx":
        text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext} (only PDF and DOCX are supported)")

    if not text:
        raise ValueError("No extractable text found in this file (it may be a scanned image).")

    return text
